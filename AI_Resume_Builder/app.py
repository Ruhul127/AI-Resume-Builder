from flask import Flask, render_template, render_template_string, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Inches
import logging
import time
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, TextAreaField
from wtforms.validators import DataRequired, Length
from openai import OpenAI

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///Resume_Builder.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Initialize OpenAI API
OPENAI_API_KEY = "your-API-key"  # Replace with your OpenAI API key
client = OpenAI(api_key=OPENAI_API_KEY)

# Configure logging
logging.basicConfig(level=logging.ERROR)

# Database Models
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    resumes = db.relationship('Resume', backref='user', lazy=True)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    job_role = db.Column(db.String(100), nullable=True)

# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    try:
        return User.query.get(int(user_id))
    except Exception as e:
        logging.error(f"Error loading user: {e}")
        return None

# Forms
class RegistrationForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(min=4, max=50)])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])

class ResumeForm(FlaskForm):
    job_role = StringField('Job Role', validators=[DataRequired()])
    content = TextAreaField('Resume Content', validators=[DataRequired()])

# Utility Functions
def generate_ai_suggestions(content, job_role):
    retries = 3  # Number of retries for the API call
    for attempt in range(retries):
        try:
            # Create a detailed prompt for OpenAI
            prompt = (
                f"Analyze the following resume content for the role of {job_role}. "
                "Provide detailed suggestions in the following format:\n\n"
                "1. **Grammar and Spelling**: Fix any errors.\n"
                "2. **Formatting**: Suggest improvements for readability.\n"
                "3. **Keywords**: Add relevant keywords from the job description.\n"
                "4. **Achievements**: Highlight quantifiable achievements.\n\n"
                f"Resume Content:\n\n{content}"
            )

            # Call the OpenAI API
            response = client.chat.completions.create(
                model="gpt-3.5-turbo", 
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that provides resume improvement suggestions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )

            # Extract and return the suggestions
            return response.choices[0].message.content.strip()
        except Exception as e:
            logging.error(f"Attempt {attempt + 1} failed: {e}")
            time.sleep(2)  # Wait before retrying

    # Fallback message if the API fails after retries
    return (
        "We couldn't generate AI suggestions at the moment. "
        "Here are some general tips to improve your resume:\n\n"
        "1. Use action verbs to describe your achievements (e.g., 'managed', 'developed', 'optimized').\n"
        "2. Tailor your resume to the job role by including relevant keywords from the job description.\n"
        "3. Keep your resume concise and avoid unnecessary details.\n"
        "4. Use bullet points for readability.\n"
        "5. Proofread for grammar and spelling errors."
    )

def generate_pdf(content):
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("Resume", styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(content, styles['BodyText']))
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error generating PDF: {e}")
        return None

def generate_docx(content):
    try:
        document = Document()
        document.add_heading('Resume', 0)
        document.add_paragraph(content)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        return None

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    if form.validate_on_submit():
        try:
            username = form.username.data
            password = generate_password_hash(form.password.data)  # Use default hashing method
            user = User(username=username, password=password)
            db.session.add(user)
            db.session.commit()
            flash('Registration successful! Please log in.')
            return redirect(url_for('login'))
        except Exception as e:
            logging.error(f"Error during registration: {e}")
            db.session.rollback()
            flash('An error occurred during registration. Please try again.')
    return render_template('register.html', form=form)

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        try:
            user = User.query.filter_by(username=form.username.data).first()
            if user and check_password_hash(user.password, form.password.data):
                login_user(user)
                return redirect(url_for('dashboard'))
            flash('Invalid username or password.')
        except Exception as e:
            logging.error(f"Error during login: {e}")
            flash('An error occurred during login. Please try again.')
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    try:
        logout_user()
    except Exception as e:
        logging.error(f"Error during logout: {e}")
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    try:
        resumes = Resume.query.filter_by(user_id=current_user.id).all()
        return render_template('dashboard.html', resumes=resumes)
    except Exception as e:
        logging.error(f"Error loading dashboard: {e}")
        flash('An error occurred while loading the dashboard. Please try again.')
        return redirect(url_for('index'))

@app.route('/resume/new', methods=['GET', 'POST'])
@login_required
def create_resume():
    form = ResumeForm()
    if form.validate_on_submit():
        try:
            content = form.content.data
            job_role = form.job_role.data
            resume = Resume(user_id=current_user.id, content=content, job_role=job_role)
            db.session.add(resume)
            db.session.commit()
            flash('Resume created successfully!')
            return redirect(url_for('dashboard'))
        except Exception as e:
            logging.error(f"Error creating resume: {e}")
            db.session.rollback()
            flash('An error occurred while creating the resume. Please try again.')
    return render_template('resume.html', form=form)

@app.route('/resume/<int:id>/delete', methods=['POST'])
@login_required
def delete_resume(id):
    try:
        # Fetch the resume from the database
        resume = Resume.query.get_or_404(id)
        
        # Ensure the user owns the resume
        if resume.user_id != current_user.id:
            flash('You do not have permission to delete this resume.')
            return redirect(url_for('dashboard'))
        
        # Delete the resume
        db.session.delete(resume)
        db.session.commit()
        flash('Resume deleted successfully!')
    except Exception as e:
        logging.error(f"Error deleting resume: {e}")
        db.session.rollback()
        flash('An error occurred while deleting the resume. Please try again.')
    return redirect(url_for('dashboard'))

@app.route('/resume/<int:id>/export/pdf')
@login_required
def export_pdf(id):
    try:
        resume = Resume.query.get_or_404(id)
        pdf_buffer = generate_pdf(resume.content)
        if pdf_buffer:
            return send_file(pdf_buffer, as_attachment=True, download_name='resume.pdf', mimetype='application/pdf')
        else:
            flash('Error generating PDF. Please try again.')
            return redirect(url_for('dashboard'))
    except Exception as e:
        logging.error(f"Error exporting PDF: {e}")
        flash('An error occurred while exporting the PDF. Please try again.')
        return redirect(url_for('dashboard'))

@app.route('/resume/<int:id>/export/docx')
@login_required
def export_docx(id):
    try:
        resume = Resume.query.get_or_404(id)
        docx_buffer = generate_docx(resume.content)
        if docx_buffer:
            return send_file(docx_buffer, as_attachment=True, download_name='resume.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            flash('Error generating Word document. Please try again.')
            return redirect(url_for('dashboard'))
    except Exception as e:
        logging.error(f"Error exporting DOCX: {e}")
        flash('An error occurred while exporting the Word document. Please try again.')
        return redirect(url_for('dashboard'))

@app.route('/resume/<int:id>/suggestions')
@login_required
def view_suggestions(id):
    try:
        # Fetch the resume from the database
        resume = Resume.query.get_or_404(id)
        
        # Ensure the user owns the resume
        if resume.user_id != current_user.id:
            flash('You do not have permission to view this resume.')
            return redirect(url_for('dashboard'))
        
        # Generate AI suggestions
        suggestions = generate_ai_suggestions(resume.content, resume.job_role)
        
        # Render the suggestions template
        return render_template('suggestions.html', suggestions=suggestions)
    
    except Exception as e:
        logging.error(f"Error generating suggestions: {e}")
        flash('An error occurred while generating suggestions. Please try again.')
        return redirect(url_for('dashboard'))

@app.route('/resume-tips')
def resume_tips():
    return render_template('resume_tips.html')

# Run the app
if __name__ == '__main__':
    try:
        with app.app_context():
            db.create_all()
        app.run(debug=True)
    except Exception as e:
        logging.error(f"Error starting the application: {e}")